import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Tiêu đề
title = doc.add_heading('BÁO CÁO CƠ SỞ Ý TƯỞNG (PAPER IDEA REPORT)', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Tên bài báo: GraphCAG: Hierarchical Cache-Augmented Generation under State Drift').bold = True
doc.add_paragraph('Người nộp/Tác giả: Nguyễn Hữu Thắng').bold = True
doc.add_paragraph('Ngày báo cáo: 15/03/2026\n')

# 1. Ý TƯỞNG CỐT LÕI
h1 = doc.add_heading('1. Ý tưởng cốt lõi (Main Idea)', level=1)
p1 = doc.add_paragraph()
p1.add_run('- Tình trạng hiện tại: ').bold = True
p1.add_run('Các hệ thống sinh tự động dựa trên học máy (LLM) phục vụ theo ngữ cảnh (như dạy học ngôn ngữ, trợ lý ảo) thường tiêu tốn nhiều chi phí tài nguyên và thời gian vì phải liên tục tái tạo lại ngữ cảnh và truy xuất dữ liệu từ đầu (bottleneck) cho từng tương tác.\n')

p1.add_run('- Hạn chế của giải pháp truyền thống: ').bold = True
p1.add_run('Hệ thống bộ nhớ đệm (Cache) khớp chính xác (Exact Cache) thường thất bại và kém an toàn khi trạng thái của hệ thống hoặc mục tiêu thay đổi (State Drift), ví dụ như khi ý định (intent) hoặc trình độ (level) của người dùng thay đổi dù truy vấn câu chữ trông rất giống nhau.\n')

p1.add_run('- Đề xuất của bài báo (GraphCAG): ').bold = True
p1.add_run('Xây dựng một kiến trúc kết hợp hệ thống sinh có augmentation và Cache (Cache-Augmented Generation - CAG) theo mô hình ')
p1.add_run('phân cấp 3 tầng (L0, L1, L2)').bold = True
p1.add_run('. Đột phá nằm ở khối trung gian ')
p1.add_run('Tầng L1 (Concept-State Near-Hit Reuse)').bold = True
p1.add_run('. GraphCAG sử dụng các phép kiểm tra tính đúng đắn dựa trên trạng thái (Profile Cache Correctness - PCC) để có thể "tái sử dụng" lại một phần cache lân cận nếu truy vấn có sai khác nhẹ mà vẫn an toàn, qua đó không cần phải gọi LLM đầy đủ mà vẫn đảm bảo được ngữ cảnh.')

# 2. QUY TRÌNH & KIẾN TRÚC
doc.add_heading('2. Quy trình & Kiến trúc (Process & Workflow)', level=1)
p2 = doc.add_paragraph('Kiến trúc của hệ thống bao gồm một bộ định tuyến đánh giá rủi ro (RAPID-GraphCAG) rẽ nhánh request thành 3 hướng chính trong một quy trình đồ thị trạng thái có hướng (LangGraph StateGraph):')

ul1 = doc.add_paragraph('Tầng L0 - Exact Reuse (Tái sử dụng khớp hoàn toàn): ', style='List Bullet')
ul1.runs[0].bold = True
ul1.add_run('Sử dụng mã băm nhanh dựa trên truy vấn chuẩn hóa và level của người dùng. Nếu có bản lưu trùng khớp hoàn toàn, trả về kết quả ngay lập tức để tiết kiệm tối đa độ trễ.\n')

ul2 = doc.add_paragraph('Tầng L1 - Concept-State Near-Hit Reuse (Tái sử dụng vùng lân cận): ', style='List Bullet')
ul2.runs[0].bold = True
ul2.add_run('Nếu L0 thất bại, thay vì phải chạy toàn bộ pipeline, hệ thống dự đoán một vân tay độ ưu tiên rẻ (Cheap Concept-State Fingerprint). Nó sẽ tìm các lịch sử tư duy/cache có chung "Khu vực đồ thị" (Graph Bucket). Nếu tính tương đồng, ý định và level (Được kiểm tra qua giao thức PCC) là chấp nhận được, hệ thống sẽ gỡ lỗi nhẹ (delta-patching) thông tin sẵn có và trả về cho người dùng.\n')

ul3 = doc.add_paragraph('Tầng L2 - Graph-Grounded Reconstruction (Tái tạo bằng Đồ thị tri thức): ', style='List Bullet')
ul3.runs[0].bold = True
ul3.add_run('Xảy ra khi cả L0 và L1 đều không an toàn để lấy cache. Hệ thống sẽ tiến hành đi hết pipeline:\n')
ul3.add_run('  + KG Expand: Mở rộng các khái niệm trên Đồ thị Knowledge Graph ưu tiên theo cấp độ (Level-Aware BFS).\n')
ul3.add_run('  + Diagnose & Route: Phân tích đánh giá (LLM sinh ra) các lỗi của câu, để chuyển hướng (Adapter, Clarify hoặc đưa thẳng qua truy xuất).\n')
ul3.add_run('  + Retrieve & Generate: Tiến hành Hybrid Retrieval tích hợp dữ liệu của Đồ Thị, vector AI và gọi qua phần sinh cuối cùng để tạo kết quả.\n')

# 3. CÁC THUẬT TOÁN VÀ CÔNG THỨC SỬ DỤNG
doc.add_heading('3. Thuật toán và Công thức toán học (Algorithms & Formulas)', level=1)

p3 = doc.add_paragraph()
p3.add_run('- Phương trình thu thập Nhận dạng Cache (Cheap Concept-State Fingerprint F_cheap):\n').bold = True
p3.add_run('Để tiết kiệm chi phí trước khi vào L2, hệ thống quét truy vấn của người dùng ra 5 thành tố:\n')
p3.add_run('F_cheap(q, P, S) = < I, l, C_seed, G_nbr, s_profile >\n')
p3.add_run('Trong đó: I là ý định sơ bộ, l là trình độ, C_seed là mầm các khái niệm chính, G_nbr là cụm nút quan hệ, và s_profile là trạng thái của phiên tương tác.\n')

p4 = doc.add_paragraph()
p4.add_run('- Công thức đánh giá Rủi ro Cache (Reuse Risk - PCC risk scoring):\n').bold = True
p4.add_run('Được sử dụng ở tầng L1, nhắm quyết định xem bản cache lịch sử có thể mang ra sử dụng hay không:\n')
p4.add_run('ρ(x, m) = 1[I(x) ≠ I(m)] + 0.5 * (1 - Jaccard(C(x), C(m))) - 0.1 * 1[Age(m) < τ_recent]\n')
p4.add_run('=> Yếu tố phạt cao nhất nếu ý định (I) hai lần sai lệch, hoặc độ bao phủ khái niệm quá xa nhau. Nếu điểm rủi ro quá lớn, chuyển về L2.\n')

p5 = doc.add_paragraph()
p5.add_run('- Thuật toán đi qua Đồ thị có Trọng số Sư phạm (PedWeight):\n').bold = True
p5.add_run('Cách đi BFS của GraphCAG sẽ thiên vị những khái niệm phù hợp với Level trình độ (l) của phiên:\n')
p5.add_run('w(r, l) = 1.0 (nếu bằng level), 0.7 (nếu cách lệch 1 level), và 0.3 (các level khác).\n')

p6 = doc.add_paragraph()
p6.add_run('- Tính điểm Truy xuất lai (Hybrid Retrieval):\n').bold = True
p6.add_run('Score(e) = α * s_kg(e) + β * s_vec(e) + γ * s_rec(e)\n')
p6.add_run('Ghép điểm của quan hệ Đồ thị tri thức, cosine tương đồng thuật toán Vector ANN, và yếu tố điểm thưởng theo thời gian học.\n')

# 4. ỨNG DỤNG THỰC TẾ & KẾT QUẢ ĐÁNH GIÁ (OVERVIEW)
doc.add_heading('4. Tổng quan kiến trúc bổ trợ & Cải thiện chi phí', level=1)
p7 = doc.add_paragraph()
p7.add_run('- Memory-Aware Model Loading: ').bold = True
p7.add_run('Hệ thống có thuật toán quản trị bộ nhớ theo budget (Priorirty Eviction) giúp lazy-load và giải phóng những phần LLM ít cần thiết.\n')
p7.add_run('- Kết quả Benchmark sơ bộ: ').bold = True
p7.add_run('Việc thêm tầng đệm ở L1 với nguyên tắc an toàn, giúp giảm thời gian chạy nền rất mạnh với hiệu năng bảo toàn cao. Ví cử như xử lý lỗi ngữ pháp (Wi+Locness) giảm từ 177.3ms xuống 100.1ms; hội thoại giảm từ 134.1ms xuống 77.7ms.')

# Save doc
doc.save('Bao_cao_y_tuong_GraphCAG.docx')
print("File Bao_cao_y_tuong_GraphCAG.docx created successfully!")
